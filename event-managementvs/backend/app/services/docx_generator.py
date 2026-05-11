import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.models.schemas import EventCircular

def create_circular_docx(data: EventCircular, output_path: str) -> str:
    """
    Creates a Word document for the event circular.
    
    Args:
        data: EventCircular object with structured content
        output_path: Path where the .docx file should be saved
        
    Returns:
        Path to the created document
    """
    doc = Document()
    
    # Add header image
    header_image_path = os.path.join(os.path.dirname(__file__), "../../assets/college_header.png")
    if os.path.exists(header_image_path):
        doc.add_picture(header_image_path, width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add spacing
    doc.add_paragraph()
    
    # Add Event Details Table
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # Table data
    table_data = [
        ("Event Title", data.event_title),
        ("Date & Time", data.date_time),
        ("Venue", data.venue),
        ("Number of Participants", data.number_of_participants),
        ("Event Type", data.event_type),
        ("Duration", data.duration)
    ]
    
    # Populate table
    for i, (field, value) in enumerate(table_data):
        row = table.rows[i]
        # Field name (bold)
        field_cell = row.cells[0]
        field_para = field_cell.paragraphs[0]
        field_run = field_para.add_run(field)
        field_run.bold = True
        
        # Field value
        value_cell = row.cells[1]
        value_cell.text = value
    
    # Add spacing after table
    doc.add_paragraph()
    
    # Event Description
    desc_heading = doc.add_paragraph()
    desc_heading_run = desc_heading.add_run("Event Description")
    desc_heading_run.bold = True
    desc_heading_run.font.size = Pt(14)
    
    doc.add_paragraph(data.event_description)
    doc.add_paragraph()
    
    # Rules
    if data.rules:
        rules_heading = doc.add_paragraph()
        rules_heading_run = rules_heading.add_run("Rules")
        rules_heading_run.bold = True
        rules_heading_run.font.size = Pt(14)
        
        for rule in data.rules:
            doc.add_paragraph(rule, style='List Bullet')
        
        doc.add_paragraph()
    
    # Judging Criteria
    if data.judging_criteria:
        criteria_heading = doc.add_paragraph()
        criteria_heading_run = criteria_heading.add_run("Judging Criteria")
        criteria_heading_run.bold = True
        criteria_heading_run.font.size = Pt(14)
        
        for criterion in data.judging_criteria:
            doc.add_paragraph(criterion, style='List Bullet')
        
        doc.add_paragraph()
    
    # Coordinators
    if data.coordinators:
        coord_heading = doc.add_paragraph()
        coord_heading_run = coord_heading.add_run("Coordinators")
        coord_heading_run.bold = True
        coord_heading_run.font.size = Pt(14)
        
        coord_para = doc.add_paragraph(", ".join(data.coordinators))
        doc.add_paragraph()
    
    # Convenor
    convenor_heading = doc.add_paragraph()
    convenor_heading_run = convenor_heading.add_run("Convenor")
    convenor_heading_run.bold = True
    convenor_heading_run.font.size = Pt(14)
    
    doc.add_paragraph(data.convenor)
    
    # Save document
    doc.save(output_path)
    return output_path
